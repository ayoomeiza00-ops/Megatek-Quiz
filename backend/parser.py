import re
import json
from docx import Document

def extract_questions_and_answers(q_docx_path, a_docx_path):
    # ----- Extract questions -----
    doc = Document(q_docx_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    text = "\n".join(full_text)

    # Regex to capture question number, stem, and options (A, B, C, D)
    pattern = r'(\d+)\.\s*(.+?)\s*(A\)\s*(.+?))\s*(B\)\s*(.+?))\s*(C\)\s*(.+?))\s*(D\)\s*(.+?))(?=\s*\d+\.|$)'
    matches = re.findall(pattern, text, re.DOTALL)

    questions = []
    for m in matches:
        q_num = int(m[0])
        stem = m[1].strip()
        opt_a = m[3].strip()
        opt_b = m[5].strip()
        opt_c = m[7].strip()
        opt_d = m[9].strip()
        questions.append({
            "id": q_num,
            "question": stem,
            "options": [opt_a, opt_b, opt_c, opt_d],
            "correct": None
        })

    # ----- Extract answers (robust) -----
    doc_ans = Document(a_docx_path)
    ans_dict = {}
    for para in doc_ans.paragraphs:
        text = para.text.strip()
        # Match any number followed by a letter (A-D) anywhere in the line
        # This handles: **1**   B\)   or 1. B) or 1 B) etc.
        match = re.search(r'(\d+)\s*[\.\)]?\s*([A-D])', text)
        if match:
            num = int(match.group(1))
            letter = match.group(2)
            ans_dict[num] = letter

    # If still empty, try a more relaxed approach: just find digits and letters separated by any non-letter
    if not ans_dict:
        for para in doc_ans.paragraphs:
            text = para.text.strip()
            # Remove all non-printable and extra spaces
            text = re.sub(r'[^A-Za-z0-9\s]', ' ', text)
            parts = text.split()
            # Look for a token that is a number (1-150) and next token is a single letter A-D
            for i, token in enumerate(parts):
                if token.isdigit() and 1 <= int(token) <= 150:
                    if i+1 < len(parts) and parts[i+1] in 'ABCD':
                        num = int(token)
                        letter = parts[i+1]
                        ans_dict[num] = letter

    # Merge answers into questions
    for q in questions:
        q_num = q["id"]
        correct_letter = ans_dict.get(q_num)
        if correct_letter:
            q["correct"] = ord(correct_letter) - 65  # A->0, B->1, C->2, D->3
        else:
            print(f"⚠️ Warning: No answer found for question {q_num}")
            q["correct"] = 0  # fallback to A

    return questions

if __name__ == "__main__":
    q_path = "Computer_Studies_Quiz_Competition_150_Questions.docx"
    a_path = "Computer_Studies_Quiz_Competition_150_Answers.docx"
    
    questions = extract_questions_and_answers(q_path, a_path)
    
    # Assign round and time limit
    for idx, q in enumerate(questions):
        q["round"] = (idx // 30) + 1
        q["time_limit"] = 15 if q["round"] <= 3 else 10

    with open("questions.json", "w") as f:
        json.dump(questions, f, indent=2)
    
    print(f"✅ questions.json created with {len(questions)} questions.")